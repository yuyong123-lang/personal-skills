"""
Merge content from a source docx (converted from markdown) into a target docx
at a specific section, with heading style remapping, table border fix, and
code block wrapping.

Usage:
    python merge_md_to_docx.py \
        --target "D:\\path\\to\\target.docx" \
        --source "/tmp/temp-content.docx" \
        --section "应用系统管理" \
        --boundary "应用商店管理" \
        --heading-shift 2 \
        --style-map '{"Heading1":"3","Heading2":"4","Heading3":"5","Heading4":"6"}'

Arguments:
    --target        Target Word document path
    --source        Source docx (converted from markdown via pandoc)
    --section       Target section heading text (insert after this)
    --boundary      Next section heading text (insert before this)
    --heading-shift Number of levels to shift (default 2, used only if --style-map not provided)
    --style-map     JSON mapping from pandoc style names to target style IDs
    --start-text    Text pattern marking the start of content to include (default: "1.")
    --fix-borders   Fix table borders in source before insertion (default: true)
    --wrap-code     Wrap code blocks in single-cell tables (default: true)
"""
import sys
import copy
import json
import argparse
from docx import Document
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def get_para_text(p_elem):
    """Get text from a paragraph XML element."""
    texts = []
    for r in p_elem.findall(qn('w:r')):
        for t in r.findall(qn('w:t')):
            texts.append(t.text or '')
    return ''.join(texts).strip()


def find_section_markers(doc, section_text, boundary_text):
    """
    Find the target section heading and boundary heading in the document.
    Returns (section_elem, boundary_elem) or raises ValueError.
    Uses parent section context to avoid matching wrong sections.
    """
    section_elem = None
    boundary_elem = None
    in_correct_context = False

    for para in doc.paragraphs:
        text = para.text.strip()
        style = para.style.name if para.style else ''

        if text == section_text:
            if style and ('heading' in style.lower() or 'Heading' in style):
                section_elem = para._element
                in_correct_context = True

        if text == boundary_text and in_correct_context:
            if style and ('heading' in style.lower() or 'Heading' in style):
                boundary_elem = para._element
                break

    return section_elem, boundary_elem


def collect_source_content(source_body, start_marker="1."):
    """
    Collect body elements from source, skipping title/TOC until start_marker.
    Returns list of XML elements.
    """
    elements = []
    skip = True

    for elem in source_body:
        if elem.tag == qn('w:sectPr'):
            continue

        if elem.tag == qn('w:p') and skip:
            text = get_para_text(elem)
            if text.startswith(start_marker):
                skip = False
            else:
                continue

        if not skip:
            elements.append(elem)

    return elements


def adjust_heading_levels(body, section_elem, boundary_elem, style_map):
    """
    Adjust heading styles for all paragraphs between section and boundary.
    Uses the provided style_map dict: {"Heading1": "3", "Heading2": "4", ...}
    """
    in_section = False
    adjusted = 0

    for elem in list(body):
        if elem is section_elem:
            in_section = True
            continue
        if elem is boundary_elem:
            in_section = False
            continue
        if not in_section or elem.tag != qn('w:p'):
            continue

        pPr = elem.find(qn('w:pPr'))
        if pPr is None:
            continue
        pStyle = pPr.find(qn('w:pStyle'))
        if pStyle is None:
            continue
        val = pStyle.get(qn('w:val'))
        if val in style_map:
            pStyle.set(qn('w:val'), style_map[val])
            adjusted += 1

    return adjusted


def fix_table_borders(doc):
    """Set full borders on all tables in the document."""
    WNS = nsdecls('w')
    border_spec = {
        'top': ('single', '4', '0', '000000'),
        'bottom': ('single', '4', '0', '000000'),
        'left': ('single', '8', '0', '000000'),
        'right': ('single', '4', '0', '000000'),
        'insideH': ('single', '4', '0', '000000'),
        'insideV': ('single', '4', '0', '000000'),
    }

    for table in doc.tables:
        tbl = table._tbl
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = parse_xml(f'<w:tblPr {WNS}></w:tblPr>')
            tbl.insert(0, tblPr)

        existing = tblPr.find(qn('w:tblBorders'))
        if existing is not None:
            tblPr.remove(existing)

        borders_xml = f'<w:tblBorders {WNS}>'
        for name, (val, sz, space, color) in border_spec.items():
            borders_xml += f'<w:{name} w:val="{val}" w:sz="{sz}" w:space="{space}" w:color="{color}"/>'
        borders_xml += '</w:tblBorders>'
        tblPr.append(parse_xml(borders_xml))


def wrap_code_blocks_in_tables(doc):
    """Find code-styled paragraphs and wrap each in a single-cell bordered table."""
    WNS = nsdecls('w')
    body = doc.element.body
    paras_to_wrap = []

    for para in body.findall(qn('w:p')):
        pPr = para.find(qn('w:pPr'))
        if pPr is not None:
            pStyle = pPr.find(qn('w:pStyle'))
            if pStyle is not None:
                val = pStyle.get(qn('w:val'))
                if val and ('code' in val.lower() or 'source' in val.lower()):
                    paras_to_wrap.append(para)

    for para in paras_to_wrap:
        parent = para.getparent()
        if parent is None:
            continue
        idx = list(parent).index(para)

        tbl = parse_xml(f'<w:tbl {WNS}></w:tbl>')
        tblPr = parse_xml(
            f'<w:tblPr {WNS}>'
            '<w:tblW w:w="0" w:type="auto"/>'
            '<w:tblBorders>'
            '<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '</w:tblBorders></w:tblPr>')
        tbl.append(tblPr)

        tr = parse_xml(f'<w:tr {WNS}></w:tr>')
        tc = parse_xml(f'<w:tc {WNS}></w:tc>')
        tcPr = parse_xml(
            f'<w:tcPr {WNS}><w:tcW w:w="0" w:type="auto"/>'
            '<w:tcBorders>'
            '<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '</w:tcBorders></w:tcPr>')
        tc.append(tcPr)
        tc.append(copy.deepcopy(para))
        tr.append(tc)
        tbl.append(tr)

        parent.remove(para)
        parent.insert(idx, tbl)


def build_default_style_map(shift):
    """
    Build default style map: Heading{i} -> str(i + shift).
    This maps to numeric style IDs that the target document uses.
    """
    return {f'Heading{i}': str(i + shift) for i in range(1, 6)}


def main():
    parser = argparse.ArgumentParser(description='Merge docx content into target section')
    parser.add_argument('--target', required=True, help='Target Word document path')
    parser.add_argument('--source', required=True, help='Source docx path (from pandoc)')
    parser.add_argument('--section', required=True, help='Target section heading text')
    parser.add_argument('--boundary', required=True, help='Next section heading text')
    parser.add_argument('--heading-shift', type=int, default=2,
                        help='Levels to shift headings (default 2, used if --style-map not provided)')
    parser.add_argument('--style-map', type=str, default=None,
                        help='JSON style mapping, e.g. \'{"Heading1":"3","Heading2":"4"}\'')
    parser.add_argument('--start-text', default='1.', help='Text marking content start in source')
    parser.add_argument('--no-fix-borders', action='store_true', help='Skip table border fix')
    parser.add_argument('--no-wrap-code', action='store_true', help='Skip code block wrapping')
    args = parser.parse_args()

    sys.stdout.reconfigure(encoding='utf-8')

    # Build style map
    if args.style_map:
        style_map = json.loads(args.style_map)
        print(f"Using custom style map: {style_map}")
    else:
        style_map = build_default_style_map(args.heading_shift)
        print(f"Using default style map (shift={args.heading_shift}): {style_map}")

    print(f"Loading source: {args.source}")
    source_doc = Document(args.source)

    # Post-process source before insertion
    if not args.no_fix_borders:
        print("Fixing table borders...")
        fix_table_borders(source_doc)

    if not args.no_wrap_code:
        print("Wrapping code blocks in tables...")
        wrap_code_blocks_in_tables(source_doc)

    # Save post-processed source (needed because we modified it)
    source_doc.save(args.source)
    source_doc = Document(args.source)

    print(f"Loading target: {args.target}")
    target_doc = Document(args.target)

    # Step 1: Find section markers
    print("\nFinding section markers...")
    section_elem, boundary_elem = find_section_markers(
        target_doc, args.section, args.boundary
    )

    if section_elem is None:
        print(f"ERROR: Could not find section '{args.section}'")
        sys.exit(1)
    print(f"  Found section: '{args.section}'")

    if boundary_elem is None:
        print(f"WARNING: Could not find boundary '{args.boundary}', appending after section")
    else:
        print(f"  Found boundary: '{args.boundary}'")

    # Step 2: Collect source content
    print("\nCollecting source content...")
    source_body = source_doc.element.body
    elements = collect_source_content(source_body, args.start_text)
    print(f"  Collected {len(elements)} elements")

    # Step 3: Insert content
    target_body = target_doc.element.body
    print("\nInserting content...")

    if boundary_elem is not None and boundary_elem in list(target_body):
        ref_idx = list(target_body).index(boundary_elem)
        for i, elem in enumerate(elements):
            new_elem = copy.deepcopy(elem)
            target_body.insert(ref_idx + i, new_elem)
        print(f"  Inserted {len(elements)} elements before boundary")
    else:
        ref_idx = list(target_body).index(section_elem)
        for elem in reversed(elements):
            new_elem = copy.deepcopy(elem)
            target_body.insert(ref_idx + 1, new_elem)
        print(f"  Inserted {len(elements)} elements after section")

    # Step 4: Adjust heading styles
    print(f"\nAdjusting heading styles with map: {style_map}...")
    adjusted = adjust_heading_levels(target_body, section_elem, boundary_elem, style_map)
    print(f"  Adjusted {adjusted} headings")

    # Step 5: Save
    print(f"\nSaving to: {args.target}")
    target_doc.save(args.target)
    print("Done!")


if __name__ == '__main__':
    main()
